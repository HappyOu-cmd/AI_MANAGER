#!/usr/bin/env python3
"""
Универсальный конвертер документов в текстовый формат
Поддерживает: PDF, DOC, DOCX, XLS, XLSX
"""

import os
import sys
from pathlib import Path
from typing import Optional
import mimetypes

from app.core.converters.base import BaseConverter
from app.utils.exceptions import DocumentConversionError


class DocumentConverter(BaseConverter):
    """Конвертирует документы различных форматов в текстовый файл"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF',
        '.doc': 'Word (DOC)',
        '.docx': 'Word (DOCX)',
        '.xls': 'Excel (XLS)',
        '.xlsx': 'Excel (XLSX)',
        '.txt': 'Text'
    }
    
    def __init__(self):
        self.detected_format = None
        
    def detect_format(self, file_path: str) -> Optional[str]:
        """Определяет формат файла по расширению и содержимому"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        # Определяем по расширению
        ext = file_path.suffix.lower()
        
        if ext in self.SUPPORTED_FORMATS:
            self.detected_format = ext
            return self.SUPPORTED_FORMATS[ext]
        
        # Пытаемся определить по MIME типу
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            if 'pdf' in mime_type:
                self.detected_format = '.pdf'
                return 'PDF'
            elif 'msword' in mime_type or 'wordprocessingml' in mime_type:
                self.detected_format = '.docx' if 'wordprocessingml' in mime_type else '.doc'
                return 'Word'
            elif 'spreadsheetml' in mime_type or 'ms-excel' in mime_type:
                self.detected_format = '.xlsx' if 'spreadsheetml' in mime_type else '.xls'
                return 'Excel'
        
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")
    
    def convert_pdf(self, file_path: str) -> str:
        """Конвертирует PDF в текст"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("Для работы с PDF установите: pip install PyPDF2")
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Страница {page_num + 1} ---\n")
                    text_content.append(text)
                    text_content.append("\n")
        
        return "\n".join(text_content)
    
    def convert_docx(self, file_path: str) -> str:
        """Конвертирует DOCX в текст с извлечением из всех элементов"""
        try:
            from docx import Document
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.run import CT_R
        except ImportError:
            raise ImportError("Для работы с DOCX установите: pip install python-docx")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            raise RuntimeError(f"Не удалось открыть DOCX файл: {e}")
        
        text_content = []
        
        # Функция для извлечения текста из элемента (включая вложенные элементы)
        def extract_text_from_element(element):
            """Рекурсивно извлекает весь текст из элемента"""
            text_parts = []
            
            # Если это параграф
            if hasattr(element, 'text'):
                if element.text and element.text.strip():
                    text_parts.append(element.text.strip())
            
            # Извлекаем текст из всех дочерних элементов
            if hasattr(element, '__iter__'):
                for child in element:
                    child_text = extract_text_from_element(child)
                    if child_text:
                        text_parts.append(child_text)
            
            # Для runs (фрагментов текста) извлекаем текст напрямую
            if hasattr(element, 'runs'):
                for run in element.runs:
                    if run.text and run.text.strip():
                        text_parts.append(run.text.strip())
            
            return ' '.join(text_parts) if text_parts else ''
        
        # Извлекаем текст из всех параграфов (включая заголовки и подзаголовки)
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_content.append(para_text)
            
            # Дополнительно извлекаем текст из runs (на случай, если paragraph.text не работает)
            if not para_text:
                runs_text = []
                for run in paragraph.runs:
                    if run.text and run.text.strip():
                        runs_text.append(run.text.strip())
                if runs_text:
                    text_content.append(' '.join(runs_text))
        
        # Извлекаем текст из таблиц (более детально)
        for table in doc.tables:
            text_content.append("\n--- Таблица ---\n")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    # Извлекаем текст из всех параграфов ячейки
                    cell_text_parts = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text_parts.append(para.text.strip())
                    # Если параграфы пустые, пробуем извлечь напрямую
                    if not cell_text_parts:
                        cell_text = cell.text.strip()
                        if cell_text:
                            cell_text_parts.append(cell_text)
                    
                    if cell_text_parts:
                        row_text.append(' '.join(cell_text_parts))
                    else:
                        row_text.append("")
                
                if any(row_text):  # Добавляем строку, если есть хотя бы одна непустая ячейка
                    text_content.append(" | ".join(row_text))
            text_content.append("\n")
        
        # Извлекаем текст из заголовков и подвалов (headers/footers)
        try:
            for section in doc.sections:
                # Заголовок
                if section.header:
                    header_texts = []
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_texts.append(para.text.strip())
                    if header_texts:
                        text_content.append(f"\n--- Заголовок ---\n")
                        text_content.extend(header_texts)
                        text_content.append("\n")
                
                # Подвал
                if section.footer:
                    footer_texts = []
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footer_texts.append(para.text.strip())
                    if footer_texts:
                        text_content.append(f"\n--- Подвал ---\n")
                        text_content.extend(footer_texts)
                        text_content.append("\n")
        except Exception as e:
            # Если не удалось извлечь заголовки/подвалы, продолжаем
            pass
        
        # Извлекаем текст из всех элементов документа (fallback метод)
        # Это может помочь, если текст находится в нестандартных местах
        try:
            # Получаем все XML элементы документа
            body = doc.element.body
            for element in body:
                if isinstance(element, CT_P):
                    # Это параграф - уже обработали выше
                    continue
                elif isinstance(element, CT_Tbl):
                    # Это таблица - уже обработали выше
                    continue
                else:
                    # Пробуем извлечь текст из других элементов
                    element_text = extract_text_from_element(element)
                    if element_text and element_text.strip():
                        text_content.append(element_text.strip())
        except Exception as e:
            # Если не удалось, продолжаем без этого метода
            pass
        
        result = "\n".join(text_content)
        
        # Если результат пустой, пробуем альтернативный метод
        if not result.strip():
            # Пробуем извлечь весь текст через XML
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Читаем основной документ
                    if 'word/document.xml' in zip_file.namelist():
                        xml_content = zip_file.read('word/document.xml')
                        root = ET.fromstring(xml_content)
                        
                        # Находим все текстовые элементы
                        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        text_elements = root.findall('.//w:t', namespace)
                        
                        if text_elements:
                            text_parts = []
                            for elem in text_elements:
                                if elem.text:
                                    text_parts.append(elem.text)
                            result = '\n'.join(text_parts)
            except Exception as e:
                pass
        
        return result
    
    def convert_doc(self, file_path: str) -> str:
        """Конвертирует DOC (старый формат Word) в текст"""
        import subprocess
        import tempfile
        from pathlib import Path
        
        # Пробуем использовать antiword (если установлен)
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except FileNotFoundError:
            pass
        except subprocess.TimeoutExpired:
            pass
        except Exception as e:
            print(f"⚠️  Antiword не сработал: {e}")
        
        # Пробуем использовать LibreOffice (если установлен)
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                # Конвертируем DOC в TXT через LibreOffice
                result = subprocess.run(
                    [
                        'libreoffice',
                        '--headless',
                        '--convert-to', 'txt',
                        '--outdir', tmpdir,
                        file_path
                    ],
                    capture_output=True,
                    timeout=60
                )
                
                if result.returncode == 0:
                    # Ищем созданный txt файл
                    input_path = Path(file_path)
                    txt_file = Path(tmpdir) / f"{input_path.stem}.txt"
                    if txt_file.exists():
                        with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                            return f.read()
        except FileNotFoundError:
            pass
        except subprocess.TimeoutExpired:
            pass
        except Exception as e:
            print(f"⚠️  LibreOffice не сработал: {e}")
        
        # Если ничего не сработало, пробуем python-docx (может не сработать для старых DOC)
        try:
            # Пытаемся открыть как DOCX (иногда работает)
            from docx import Document
            try:
                doc = Document(file_path)
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                if text_content:
                    return "\n".join(text_content)
            except:
                pass
        except ImportError:
            pass
        
        # Если все методы не сработали
        raise RuntimeError(
            "Не удалось конвертировать DOC файл.\n"
            "Установите один из инструментов:\n"
            "  1. antiword: sudo apt-get install antiword\n"
            "  2. LibreOffice: sudo apt-get install libreoffice\n"
            "Или откройте файл в Word и сохраните как DOCX"
        )
    
    def convert_xlsx(self, file_path: str) -> str:
        """Конвертирует XLSX в текст"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("Для работы с Excel установите: pip install pandas openpyxl")
        
        text_content = []
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            text_content.append(f"\n{'='*60}")
            text_content.append(f"Лист: {sheet_name}")
            text_content.append(f"{'='*60}\n")
            
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Заголовки
            if not df.empty:
                headers = df.columns.tolist()
                text_content.append(" | ".join(str(h) for h in headers))
                text_content.append("-" * 60)
                
                # Данные
                for idx, row in df.iterrows():
                    row_data = []
                    for col in headers:
                        value = row[col]
                        if pd.notna(value):
                            row_data.append(str(value))
                        else:
                            row_data.append("")
                    text_content.append(" | ".join(row_data))
            
            text_content.append("\n")
        
        return "\n".join(text_content)
    
    def convert_xls(self, file_path: str) -> str:
        """Конвертирует XLS (старый формат Excel) в текст"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("Для работы с Excel установите: pip install pandas xlrd")
        
        # XLS требует xlrd
        try:
            import xlrd
        except ImportError:
            raise ImportError("Для работы с XLS установите: pip install xlrd")
        
        return self.convert_xlsx(file_path)  # pandas может работать с XLS через xlrd
    
    def convert_txt(self, file_path: str) -> str:
        """Просто читает текстовый файл"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def convert(self, input_file: str, output_file: Optional[str] = None) -> str:
        """
        Конвертирует документ в текстовый файл
        
        Args:
            input_file: Путь к входному файлу
            output_file: Путь к выходному файлу (если None, создается автоматически)
        
        Returns:
            Путь к созданному текстовому файлу
        """
        input_path = Path(input_file)
        
        # Определяем формат
        format_name = self.detect_format(str(input_path))
        print(f"📄 Обнаружен формат: {format_name}")
        
        # Выбираем метод конвертации
        ext = self.detected_format
        
        if ext == '.pdf':
            text = self.convert_pdf(str(input_path))
        elif ext == '.docx':
            text = self.convert_docx(str(input_path))
        elif ext == '.doc':
            text = self.convert_doc(str(input_path))
        elif ext == '.xlsx':
            text = self.convert_xlsx(str(input_path))
        elif ext == '.xls':
            text = self.convert_xls(str(input_path))
        elif ext == '.txt':
            text = self.convert_txt(str(input_path))
        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")
        
        # Определяем имя выходного файла
        if output_file is None:
            output_path = input_path.parent / f"{input_path.stem}_converted.txt"
        else:
            output_path = Path(output_file)
        
        # Сохраняем текст
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        print(f"✅ Конвертировано в: {output_path}")
        print(f"📊 Размер: {len(text)} символов, {len(text.splitlines())} строк")
        
        return str(output_path)


def main():
    """Главная функция для запуска из командной строки"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Конвертер документов в текстовый формат',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Поддерживаемые форматы:
  - PDF (.pdf)
  - Word (.doc, .docx)
  - Excel (.xls, .xlsx)
  - Text (.txt)

Примеры использования:
  python document_converter.py document.pdf
  python document_converter.py report.docx -o output.txt
  python document_converter.py data.xlsx
        """
    )
    
    parser.add_argument('input_file', help='Путь к входному файлу')
    parser.add_argument('-o', '--output', help='Путь к выходному файлу (по умолчанию: имя_файла_converted.txt)')
    
    args = parser.parse_args()
    
    try:
        converter = DocumentConverter()
        output_path = converter.convert(args.input_file, args.output)
        print(f"\n✅ Готово! Текст сохранен в: {output_path}")
    except Exception as e:
        print(f"\n❌ Ошибка: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

