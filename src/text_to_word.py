#!/usr/bin/env python3
"""
Модуль для конвертации текста в Word документ
"""

from pathlib import Path
from typing import Optional
import re


class TextToWordConverter:
    """Конвертирует текст в Word документ"""
    
    def __init__(self):
        """Инициализация конвертера"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self.Document = Document
            self.Pt = Pt
            self.Inches = Inches
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            raise ImportError("Для работы с Word установите: pip install python-docx")
    
    def convert(self, text: str, output_path: str, title: Optional[str] = None) -> str:
        """
        Создаёт Word документ из текста
        
        Args:
            text: Текст для вставки
            output_path: Путь к выходному файлу
            title: Заголовок документа (опционально)
        
        Returns:
            Путь к созданному файлу
        """
        if not self.docx_available:
            raise ImportError("python-docx не установлен")
        
        doc = self.Document()
        
        # Добавляем заголовок
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        
        # Парсим текст
        if self._is_table(text):
            self._add_table_from_text(doc, text)
        else:
            # Разбиваем на параграфы
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    # Проверяем, не заголовок ли это (начинается с #)
                    if line.startswith('#'):
                        level = len(line) - len(line.lstrip('#'))
                        level = min(level, 6)  # Максимум 6 уровней
                        doc.add_heading(line.lstrip('# ').strip(), level=level)
                    else:
                        para = doc.add_paragraph(line)
                        # Форматируем жирный текст (**текст**)
                        self._format_markdown(para, line)
        
        # Сохраняем
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path_obj))
        
        return str(output_path_obj)
    
    def _is_table(self, text: str) -> bool:
        """Проверяет, содержит ли текст таблицу (Markdown формат)"""
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if len(lines) < 2:
            return False
        
        # Проверяем наличие разделителя таблицы
        has_separator = any('|' in line and all(c in '-|: ' for c in line) for line in lines)
        has_pipes = any('|' in line for line in lines[:3])  # Проверяем первые 3 строки
        
        return has_pipes and has_separator
    
    def _add_table_from_text(self, doc, text: str):
        """Добавляет таблицу из Markdown текста"""
        lines = [l for l in text.split('\n') if l.strip()]
        
        # Находим заголовок таблицы
        header_line = None
        separator_line = None
        data_start = 0
        
        for i, line in enumerate(lines):
            if '|' in line:
                if header_line is None:
                    header_line = line
                elif separator_line is None and all(c in '-|: ' for c in line):
                    separator_line = line
                    data_start = i + 1
                    break
        
        if not header_line:
            # Если не таблица, просто добавляем как текст
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
            return
        
        # Парсим заголовок
        headers = [h.strip() for h in header_line.split('|') if h.strip() or h == '']
        headers = [h for h in headers if h]  # Убираем пустые
        
        if not headers:
            return
        
        # Создаём таблицу
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Заголовок
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(header_cells):
                header_cells[i].text = header
                if header_cells[i].paragraphs:
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Данные
        for line in lines[data_start:]:
            if '|' in line and not all(c in '-|: ' for c in line):
                row_cells = table.add_row().cells
                values = [v.strip() for v in line.split('|')]
                # Убираем пустые значения с краёв
                while values and not values[0]:
                    values.pop(0)
                while values and not values[-1]:
                    values.pop()
                
                for i, value in enumerate(values[:len(headers)]):
                    if i < len(row_cells):
                        row_cells[i].text = value
    
    def _format_markdown(self, para, text: str):
        """Форматирует Markdown в параграфе (жирный текст)"""
        # Простая обработка **текст**
        pattern = r'\*\*(.*?)\*\*'
        parts = re.split(pattern, text)
        
        run = para.runs[0] if para.runs else para.add_run()
        run.text = ''
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Обычный текст
                if part:
                    run = para.add_run(part)
            else:
                # Жирный текст
                run = para.add_run(part)
                run.bold = True

